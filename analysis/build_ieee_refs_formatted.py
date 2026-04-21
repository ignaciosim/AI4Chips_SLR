#!/usr/bin/env python3
"""
Produce two formatted deliverables for the 52 shortlist papers:

  scopus_out10/references_ieee.docx  — Word document with hyperlinked DOIs,
                                        italic venue names, hanging indent.
                                        Paste into manuscript's References.

  scopus_out10/references.bib         — BibTeX file for import into
                                        Zotero/Mendeley/EndNote.

Numbering matches scopus_out10/references_ieee.md ([10]–[61] by default).

Reads:
  scopus_out10/final_ai4chips_high_only.json
  scopus_out10/openalex_cache/openalex_ai4chips.json
  scopus_out10/stage_shortlists.md
"""
from __future__ import annotations

import argparse
import json
import re
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Reuse helpers from build_ieee_refs.py (same directory)
import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_ieee_refs import (
    VENUE_ABBREV, MONTH, ieee_name, format_authors, abbrev_venue,
    month_from_date, load_shortlist_dois,
)


# ── DOCX hyperlink helper ──────────────────────────────────────────────────

HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"


def add_hyperlink(paragraph, url, text, color="0000EE"):
    """Insert a clickable hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, HYPERLINK_REL, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def write_ieee_paragraph(doc, n, paper, authors):
    """Write one reference entry as a hanging-indent paragraph with italic
    venue and a live DOI hyperlink."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.left_indent = Inches(0.35)
    pf.first_line_indent = Inches(-0.35)
    pf.space_after = Pt(4)

    # Plain text: "[N] Authors, “Title,” "
    title = paper["title"].strip().rstrip(".")
    author_str = format_authors(authors) or (paper.get("creator") or "Unknown")
    prefix_run = para.add_run(f"[{n}] {author_str}, \u201C{title},\u201D ")

    # Italic venue
    venue = abbrev_venue(paper.get("publication") or "")
    if venue:
        venue_run = para.add_run(venue)
        venue_run.italic = True

    # Trailing metadata
    vol = paper.get("volume") or ""
    issue = paper.get("issue") or ""
    pages = paper.get("pages") or ""
    year = paper["year"]
    month = month_from_date(paper.get("cover_date") or "")
    trailing = []
    if vol:
        trailing.append(f"vol. {vol}")
    if issue:
        trailing.append(f"no. {issue}")
    if pages:
        trailing.append(f"pp. {pages}")
    if month:
        trailing.append(f"{month} {year}")
    elif year:
        trailing.append(str(year))
    if trailing:
        para.add_run(", " + ", ".join(trailing))

    # DOI as hyperlink
    doi = (paper.get("doi") or "").strip()
    if doi:
        para.add_run(", doi: ")
        add_hyperlink(para, f"https://doi.org/{doi}", doi)
    para.add_run(".")


# ── BibTeX ─────────────────────────────────────────────────────────────────

def bibtex_key(paper, authors, existing_keys: set[str]) -> str:
    """Return a unique BibTeX key of the form 'SurnameYYYY' with a letter
    suffix on collision."""
    if authors:
        surname = authors[0].split()[-1]
    else:
        creator = paper.get("creator") or "Unknown"
        surname = re.split(r"[ ,]", creator.strip())[0]
    # Strip non-ASCII letters for clean BibTeX keys
    surname = re.sub(r"[^A-Za-z]", "", surname) or "Anon"
    base = f"{surname}{paper['year']}"
    key = base
    suffix = 0
    while key in existing_keys:
        suffix += 1
        key = f"{base}{chr(ord('a') + suffix - 1)}"
    existing_keys.add(key)
    return key


def bibtex_name(full_name: str) -> str:
    """'Brendan Dolan-Gavitt' → 'Dolan-Gavitt, Brendan'.
    'Cheng K.C.C.' (Scopus-style) → 'Cheng, K.C.C.'"""
    full_name = full_name.strip()
    if not full_name:
        return ""
    parts = full_name.split()
    if len(parts) == 1:
        return parts[0]
    # Scopus creator format: 'Surname Initials' where initials end with '.'
    if parts[-1].endswith("."):
        return f"{parts[0]}, {' '.join(parts[1:])}"
    # OpenAlex format: 'First [Middle...] Last'
    return f"{parts[-1]}, {' '.join(parts[:-1])}"


def bibtex_entry(key: str, paper, authors) -> str:
    author_parts = [bibtex_name(a) for a in authors if a.strip()]
    if not author_parts:
        author_parts = [bibtex_name(paper.get("creator") or "Unknown")]
    author_field = " and ".join(author_parts)
    title = paper["title"].strip().rstrip(".")
    # Preserve title case by double-braces
    title_field = "{{" + title + "}}"
    journal = paper.get("publication") or ""
    vol = paper.get("volume") or ""
    issue = paper.get("issue") or ""
    pages = paper.get("pages") or ""
    if pages and "-" in pages:
        # Canonical BibTeX double-hyphen for page ranges
        pages = pages.replace("-", "--")
    year = paper["year"]
    cover_date = paper.get("cover_date") or ""
    m = re.match(r"\d{4}-(\d{2})", cover_date)
    month_abbr = {
        "01": "jan", "02": "feb", "03": "mar", "04": "apr", "05": "may",
        "06": "jun", "07": "jul", "08": "aug", "09": "sep", "10": "oct",
        "11": "nov", "12": "dec",
    }.get(m.group(1), "") if m else ""
    doi = (paper.get("doi") or "").strip()

    fields = []
    fields.append(f"  author    = {{{author_field}}}")
    fields.append(f"  title     = {title_field}")
    if journal:
        fields.append(f"  journal   = {{{journal}}}")
    if vol:
        fields.append(f"  volume    = {{{vol}}}")
    if issue:
        fields.append(f"  number    = {{{issue}}}")
    if pages:
        fields.append(f"  pages     = {{{pages}}}")
    fields.append(f"  year      = {{{year}}}")
    if month_abbr:
        fields.append(f"  month     = {month_abbr}")
    if doi:
        fields.append(f"  doi       = {{{doi}}}")
    body = ",\n".join(fields)
    return f"@article{{{key},\n{body}\n}}"


# ── Main ───────────────────────────────────────────────────────────────────

def write_existing_entry(doc, entry):
    """Write a pre-existing reference (numbered by the author) as a hanging-
    indent paragraph with italic venue and live DOI hyperlink. Handles
    placeholder slots by emitting a short '[Placeholder]' line so the number
    is still visible in the list."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.left_indent = Inches(0.35)
    pf.first_line_indent = Inches(-0.35)
    pf.space_after = Pt(4)
    n = entry["n"]
    if entry.get("placeholder"):
        run = para.add_run(f"[{n}] {entry.get('hint', '[Placeholder]')}")
        run.italic = True
        return
    # IEEE book style: authors, *Title*. Publisher, Year.
    # IEEE article/proc. style: authors, "Title," *Venue*, vol..., year.
    if entry.get("kind") == "book":
        para.add_run(f"[{n}] {entry['authors']}, ")
        tr = para.add_run(entry["title"])
        tr.italic = True
        para.add_run(". ")
        if entry.get("trailing"):
            para.add_run(entry["trailing"])
    else:
        para.add_run(f"[{n}] {entry['authors']}, \u201C{entry['title']},\u201D ")
        if entry.get("venue"):
            v = para.add_run(entry["venue"])
            v.italic = True
        if entry.get("trailing"):
            para.add_run(f", {entry['trailing']}")
    if entry.get("doi"):
        para.add_run(", doi: ")
        add_hyperlink(para, f"https://doi.org/{entry['doi']}", entry["doi"])
    para.add_run(".")


def existing_bibtex(entry: dict) -> str:
    """Build a BibTeX entry for an existing (author-supplied) reference. Skips
    placeholder slots (returns None)."""
    if entry.get("placeholder"):
        return None
    # Parse authors into BibTeX format (Last, First and Last, First)
    authors_raw = entry["authors"]
    # Split on ", " and " and ", handling the final " and "
    raw = re.split(r",\s+and\s+|\s+and\s+|,\s+", authors_raw)
    bib_authors = []
    for a in raw:
        a = a.strip().rstrip(",")
        if not a:
            continue
        # Already in "F. G. Lastname" form → convert to "Lastname, F. G."
        parts = a.split()
        if len(parts) >= 2:
            # Combine multi-word surnames (e.g., Dolan-Gavitt): surname is last token
            surname = parts[-1]
            given = " ".join(parts[:-1])
            bib_authors.append(f"{surname}, {given}")
        else:
            bib_authors.append(a)
    author_field = " and ".join(bib_authors)
    title_field = "{{" + entry["title"] + "}}"
    key = entry["bibtex_key"]
    kind = entry.get("bibtex_type", "article")
    extra = entry.get("bibtex_extra", {})
    # Extract year from trailing
    year_match = re.search(r"(19|20)\d{2}", entry.get("trailing") or "")
    year = year_match.group(0) if year_match else ""
    fields = [f"  author    = {{{author_field}}}", f"  title     = {title_field}"]
    for k, v in extra.items():
        fields.append(f"  {k:9s} = {{{v}}}")
    if year:
        fields.append(f"  year      = {{{year}}}")
    if entry.get("doi"):
        fields.append(f"  doi       = {{{entry['doi']}}}")
    return f"@{kind}{{{key},\n" + ",\n".join(fields) + "\n}"


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--datadir", default="scopus_out10")
    ap.add_argument("--start", type=int, default=15,
                    help="First reference number for shortlist papers "
                         "(default 15 = after existing [1]\u2013[14])")
    args = ap.parse_args()

    outdir = Path(args.datadir)
    scopus_data = {
        (p.get("doi") or "").lower(): p
        for p in json.loads((outdir / "final_ai4chips_high_only.json").read_text())
    }
    oa_cache = json.loads((outdir / "openalex_cache" / "openalex_ai4chips.json").read_text())
    author_lookup = {
        doi.lower(): list((info.get("author_names") or {}).values())
        for doi, info in oa_cache.items()
    }
    shortlist = load_shortlist_dois(outdir / "stage_shortlists.md")
    ordered = list(OrderedDict(
        (doi, (stage, role)) for doi, stage, role in shortlist
    ).items())

    existing_refs_path = outdir / "existing_refs.json"
    existing_refs = json.loads(existing_refs_path.read_text()) if existing_refs_path.exists() else []

    # ── DOCX output ──────────────────────────────────────────────────────
    doc = Document()
    heading = doc.add_paragraph()
    heading_run = heading.add_run("References")
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    note = doc.add_paragraph(
        f"[1]\u2013[{len(existing_refs)}]: pre-existing references in the manuscript "
        f"(placeholder slots marked where empty). "
        f"[{args.start}]\u2013[{args.start + len(ordered) - 1}]: shortlist papers. "
        "DOIs are live hyperlinks. Journal names are italicized."
    )
    note.paragraph_format.space_after = Pt(12)

    # Existing author-supplied references [1]–[14]
    for entry in existing_refs:
        write_existing_entry(doc, entry)
    # Separator
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(8)
    sep_run = sep.add_run("--- shortlist papers below ---")
    sep_run.italic = True
    sep_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Shortlist papers
    for i, (doi, _) in enumerate(ordered, start=args.start):
        p = scopus_data.get(doi)
        if not p:
            doc.add_paragraph(f"[{i}] DOI {doi} not found in Scopus metadata.")
            continue
        authors = author_lookup.get(doi, [])
        write_ieee_paragraph(doc, i, p, authors)
    docx_path = outdir / "references_ieee.docx"
    doc.save(docx_path)
    print(f"Wrote {docx_path}")

    # ── BibTeX output ────────────────────────────────────────────────────
    bib_lines = [
        "% References for AI-for-Chips SLR.",
        f"% Slots [1]\u2013[{len(existing_refs)}]: pre-existing references from the manuscript.",
        f"% Slots [{args.start}]\u2013[{args.start + len(ordered) - 1}]: shortlist papers.",
        "% Import into Zotero / Mendeley / EndNote via File \u2192 Import.",
        "% BibTeX keys use FirstAuthorSurname+Year; reference managers renumber",
        "% automatically based on citation order in your document.",
        "",
        "% ---- pre-existing references (author-supplied) ----",
        "",
    ]
    for entry in existing_refs:
        bib = existing_bibtex(entry)
        if bib is None:
            bib_lines.append(f"% [{entry['n']}] placeholder \u2014 add BibTeX entry when the reference is known.")
            bib_lines.append("")
        else:
            bib_lines.append(bib)
            bib_lines.append("")

    bib_lines.append("% ---- shortlist papers ----")
    bib_lines.append("")
    keys_seen: set[str] = set(e["bibtex_key"] for e in existing_refs if "bibtex_key" in e)
    for doi, _ in ordered:
        p = scopus_data.get(doi)
        if not p:
            bib_lines.append(f"% DOI {doi}: not found in Scopus metadata.\n")
            continue
        authors = author_lookup.get(doi, [])
        key = bibtex_key(p, authors, keys_seen)
        bib_lines.append(bibtex_entry(key, p, authors))
        bib_lines.append("")
    bib_path = outdir / "references.bib"
    bib_path.write_text("\n".join(bib_lines))
    print(f"Wrote {bib_path}")

    print(f"\n{len(existing_refs)} pre-existing references \u2192 [1]\u2013[{len(existing_refs)}] "
          f"({sum(1 for e in existing_refs if e.get('placeholder'))} placeholders)")
    print(f"{len(ordered)} shortlist papers \u2192 [{args.start}]\u2013[{args.start + len(ordered) - 1}]")


if __name__ == "__main__":
    main()
